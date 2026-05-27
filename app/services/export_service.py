"""Export service for generating resumes and cover letters in multiple formats."""

import io

from loguru import logger


def export_to_markdown(content: str, title: str = "Resume") -> str:
    """Content is already markdown — just return it."""
    return content


def export_to_html(content: str, title: str = "Resume") -> str:
    """Convert markdown content to styled HTML."""
    try:
        import markdown

        html_body = markdown.markdown(content, extensions=["tables", "fenced_code"])
    except ImportError:
        # Fallback: simple conversion
        html_body = content.replace("\n", "<br>\n")

    return f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        body {{
            font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
            max-width: 800px; margin: 0 auto;
            padding: 40px 20px; line-height: 1.6; color: #333;
        }}
        h1 {{ color: #2c3e50; border-bottom: 2px solid #3498db; padding-bottom: 10px; }}
        h2 {{ color: #2980b9; margin-top: 30px; }}
        h3 {{ color: #7f8c8d; }}
        ul {{ padding-left: 20px; }}
        li {{ margin-bottom: 5px; }}
        a {{ color: #3498db; }}
        .header {{ text-align: center; margin-bottom: 30px; }}
    </style>
</head>
<body>
    {html_body}
</body>
</html>"""


def export_to_docx(content: str, title: str = "Resume") -> bytes:
    """Convert markdown content to DOCX."""
    try:
        from docx import Document
        from docx.shared import Pt

        doc = Document()

        # Set default font
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)

        # Parse markdown content line by line
        lines = content.split("\n")
        for line in lines:
            line = line.strip()
            if not line:
                continue
            elif line.startswith("# "):
                doc.add_heading(line[2:], level=1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif line.startswith("**") and line.endswith("**"):
                p = doc.add_paragraph()
                run = p.add_run(line.strip("*"))
                run.bold = True
            else:
                doc.add_paragraph(line)

        buffer = io.BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
    except Exception as e:
        logger.error(f"DOCX export failed: {e}")
        return b""


def export_to_pdf(content: str, title: str = "Resume") -> bytes:
    """Convert markdown content to PDF using fpdf2 (pure Python, no system deps)."""
    try:
        from fpdf import FPDF

        class ResumePDF(FPDF):
            def header(self):
                self.set_font("Helvetica", "B", 16)
                self.cell(0, 10, title, align="C", new_x="LMARGIN", new_y="NEXT")
                self.ln(4)

            def footer(self):
                self.set_y(-15)
                self.set_font("Helvetica", "I", 8)
                self.cell(0, 10, f"Page {self.page_no()}/{{nb}}", align="C")

        pdf = ResumePDF()
        pdf.alias_nb_pages()
        pdf.set_auto_page_break(auto=True, margin=20)
        pdf.add_page()

        lines = content.split("\n")
        for line in lines:
            stripped = line.strip()
            if not stripped:
                pdf.ln(3)
                continue

            if stripped.startswith("# "):
                pdf.set_font("Helvetica", "B", 14)
                pdf.multi_cell(0, 7, stripped[2:])
                pdf.ln(2)
            elif stripped.startswith("## "):
                pdf.set_font("Helvetica", "B", 12)
                pdf.multi_cell(0, 6, stripped[3:])
                pdf.ln(1)
            elif stripped.startswith("### "):
                pdf.set_font("Helvetica", "B", 11)
                pdf.multi_cell(0, 6, stripped[4:])
                pdf.ln(1)
            elif stripped.startswith("- ") or stripped.startswith("* "):
                pdf.set_font("Helvetica", "", 10)
                pdf.cell(8)
                pdf.multi_cell(0, 5, f"• {stripped[2:]}")
            elif stripped.startswith("**") and stripped.endswith("**"):
                pdf.set_font("Helvetica", "B", 10)
                pdf.multi_cell(0, 5, stripped.strip("*"))
            else:
                pdf.set_font("Helvetica", "", 10)
                pdf.multi_cell(0, 5, stripped)

        return bytes(pdf.output())
    except Exception as e:
        logger.error(f"PDF export failed: {e}")
        return b""


def export_to_latex(content: str, title: str = "Resume", resume=None) -> str:
    """Export to LaTeX (via template engine) or fallback to markdown conversion."""
    if resume is not None:
        from app.services.template_engine import render_resume

        return render_resume(resume)

    # Fallback: convert markdown to LaTeX
    lines = content.split("\n")
    latex_lines = [
        r"\documentclass[11pt,a4paper]{article}",
        r"\usepackage[utf8]{inputenc}",
        r"\usepackage[margin=1in]{geometry}",
        r"\usepackage{enumitem}",
        r"\usepackage{hyperref}",
        r"\usepackage{titlesec}",
        r"",
        r"\titleformat{\section}{\large\bfseries}{}{0em}{}[\titlerule]",
        r"\titlespacing{\section}{0pt}{12pt}{6pt}",
        r"",
        r"\begin{document}",
        "",
    ]

    in_itemize = False
    for line in lines:
        line = line.strip()
        if not line:
            if in_itemize:
                latex_lines.append(r"\end{itemize}")
                in_itemize = False
            latex_lines.append("")
        elif line.startswith("# "):
            if in_itemize:
                latex_lines.append(r"\end{itemize}")
                in_itemize = False
            latex_lines.append(r"\begin{center}")
            latex_lines.append(r"{\LARGE\bfseries " + _escape_latex_chars(line[2:]) + r"}")
            latex_lines.append(r"\end{center}")
        elif line.startswith("## "):
            if in_itemize:
                latex_lines.append(r"\end{itemize}")
                in_itemize = False
            latex_lines.append(r"\section{" + _escape_latex_chars(line[3:]) + "}")
        elif line.startswith("### "):
            if in_itemize:
                latex_lines.append(r"\end{itemize}")
                in_itemize = False
            latex_lines.append(r"\subsection{" + _escape_latex_chars(line[4:]) + "}")
        elif line.startswith("- ") or line.startswith("* "):
            if not in_itemize:
                latex_lines.append(r"\begin{itemize}[leftmargin=*]")
                in_itemize = True
            latex_lines.append(r"  \item " + _escape_latex_chars(line[2:]))
        elif line.startswith("**") and line.endswith("**"):
            if in_itemize:
                latex_lines.append(r"\end{itemize}")
                in_itemize = False
            latex_lines.append(r"\textbf{" + _escape_latex_chars(line.strip("*")) + "}")
        else:
            if in_itemize:
                latex_lines.append(r"\end{itemize}")
                in_itemize = False
            latex_lines.append(_escape_latex_chars(line))

    if in_itemize:
        latex_lines.append(r"\end{itemize}")

    latex_lines.append(r"\end{document}")
    return "\n".join(latex_lines)


def _escape_latex_chars(text: str) -> str:
    """Escape special LaTeX characters."""
    for char in ["\\", "&", "%", "$", "#", "_", "~", "^"]:
        if char == "\\":
            text = text.replace(char, r"\textbackslash{}")
        elif char == "~":
            text = text.replace(char, r"\textasciitilde{}")
        elif char == "^":
            text = text.replace(char, r"\textasciicircum{}")
        else:
            text = text.replace(char, "\\" + char)
    return text


def export_content(
    content: str, format_type: str, title: str = "Resume", resume=None
) -> tuple[bytes | str, str]:
    """Export content to the specified format.

    Returns:
        Tuple of (content, media_type). content is bytes for binary formats, str for text.
    """
    exporters = {
        "text": (export_to_markdown, "text/plain"),
        "markdown": (export_to_markdown, "text/plain"),
        "html": (export_to_html, "text/html"),
        "docx": (
            export_to_docx,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ),
        "pdf": (export_to_pdf, "application/pdf"),
        "latex": (export_to_latex, "text/plain"),
        "tex": (export_to_latex, "text/plain"),
    }

    key = format_type.lower()
    if key not in exporters:
        key = "text"

    exporter, media_type = exporters[key]
    if key in ("latex", "tex") and resume is not None:
        result = exporter(content, title, resume=resume)
    else:
        result = exporter(content, title)

    if isinstance(result, bytes):
        return result, media_type
    return result.encode("utf-8"), media_type
